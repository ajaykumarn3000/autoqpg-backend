from docx import Document
import zipfile
# import os
import base64


def read_docs(file_path: str):
    """Reads .docx file from the file_path"""

    # Load the .docx file
    document = Document(file_path)

    # Gets the table from the document (Assuming all data present in the first table)
    table = document.tables[0]
    qb_data = []
    image_generator = extract_images(file_path)

    for row_index, row in enumerate(table.rows):
        question_data = {"para": [], "image": []}
        try:
            question_data["serial_no"] = int(row.cells[0].text.strip())
            question_data["co"] = int(row.cells[2].text.strip())
        except ValueError:
            return f"Error: Serial Number and Module Number must be integers. Check row {row_index + 1}"
        for paragraph in row.cells[1].paragraphs:
            question_text = paragraph.text.strip()
            if question_text:
                question_data["para"].append(question_text)
            if paragraph.runs:
                for run in paragraph.runs:
                    if run.element.xpath(".//a:blip"):
                        question_data["image"].append(next(image_generator))
        qb_data.append(question_data)
    
    return qb_data


def extract_images(docx_path):
    with zipfile.ZipFile(docx_path, "r") as docx_file:
        for file in docx_file.namelist():
            if file.startswith("word/media/"):
                image_data = docx_file.read(file)
                base64_image = base64.b64encode(image_data).decode("utf-8")
                image_url = f"data:image/png;base64,{base64_image}"
                yield image_url
               

read_docs("temp/testing.docx")
