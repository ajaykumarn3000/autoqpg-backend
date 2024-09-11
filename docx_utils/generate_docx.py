from docx import Document
from docx.shared import Inches, Cm

from docx_utils.styles import add_styles

CHAR_INDEX = "abcdefghijklmnopqrstuvwxyz"
CAPITAL_CHAR_INDEX = "ABCDEFGHIJKLMNOPQRSTUVWXYZ"

def generate_docx(data, save_path):
    # Create a new Document
    doc = Document("docx_utils/question_paper_template.docx")
    add_styles(doc)

    # Add Title
    title = doc.add_paragraph(data["title"], "title")
    acad_year = doc.add_paragraph(f"Academic Year: {data['acad_year']}", "acad_year")

    # Add Department, Year, Course table
    table = doc.add_table(rows=5, cols=2)
    table.style = 'table'

    table.cell(0, 0).text = f"Department: {data['department']}"
    table.cell(0, 1).text = f"Year: {data['year']}\tSemester: {data['semester']}"

    table.cell(1, 0).text = f"Course: {data['course']}"
    table.cell(1, 1).text = f"Time: {data['time']}"

    table.cell(2, 0).text = f"Date: {data['date']}"
    table.cell(2, 1).text = f"No. of Pages: {data['total_pages']}"

    table.cell(3, 0).text = f"Marks: {data['total_marks']}"

    # Set column width
    for row in table.rows:
        row.cells[0].width = Cm(12.4)  # First column width
        row.cells[1].width = Cm(5)

    # Add Instructions
    instructions = doc.add_paragraph()
    instructions.add_run("Instructions: ", "instruction_title")
    instructions.add_run(data["instructions"], "instruction_body")

    # Add Notes
    notes = doc.add_paragraph()
    notes.add_run("Note:", "note_title")
    for n, note in enumerate(data["note"]):
        notes.add_run(f"\n{n+1}. {note}", "note_body")

    ques_table = doc.add_table(rows=1, cols=6)
    ques_table.style = "Table Grid"

    # Add questions table header
    header_cells = ["Marks", "CO", "BL", "DL"]
    for i, text in enumerate(header_cells, start=2):
        para = ques_table.cell(0, i).paragraphs[0]
        para.style = "ques_table_head"
        para.text = text

    # Add questions
    for ques_no, question in enumerate(data["questions"]):
        row = ques_table.add_row()
        qno = row.cells[0].paragraphs[0]
        qno.style = "ques_table_head"
        qno.text = f"Q.{ques_no+1}"

        ques = row.cells[1].paragraphs[0]
        ques.style = "ques_table_body"
        ques.text = "Answer the following questions:"

        # Add sub questions
        for sub_ques_no, sub_question in enumerate(question):
            sub_row = ques_table.add_row()
            sub_qno = sub_row.cells[0].paragraphs[0]
            sub_qno.text = f"{CHAR_INDEX[sub_ques_no]}."
            sub_qno.style = "ques_table_body_center"

            sub_ques = sub_row.cells[1].paragraphs[0]
            sub_ques.text = "\n".join(sub_question["question"])
            sub_ques.style = "ques_table_body"

            for image in sub_question["image"]:
                run = sub_ques.add_run()
                run.add_picture(image, width=Inches(1.5))

            marks = sub_row.cells[2].paragraphs[0]
            marks.text = str(sub_question["marks"])
            marks.style = "ques_table_body_center"

            co = sub_row.cells[3].paragraphs[0]
            co.text = f"CO{str(sub_question['co'])}"
            co.style = "ques_table_body_center"

            bl = sub_row.cells[4].paragraphs[0]
            bl.text = f"BL{str(sub_question['bl'])}"
            bl.style = "ques_table_body_center"

            dl = sub_row.cells[5].paragraphs[0]
            dl.text = sub_question["dl"]
            dl.style = "ques_table_body_center"

    # Set column width
    for row in ques_table.rows:
        row.cells[0].width = Cm(1.2)
        row.cells[1].width = Cm(15)
        row.cells[2].width = Cm(1.2)
        row.cells[3].width = Cm(1.2)
        row.cells[4].width = Cm(1.2)
        row.cells[5].width = Cm(1.2)

    # Save the document
    doc.save(save_path)
