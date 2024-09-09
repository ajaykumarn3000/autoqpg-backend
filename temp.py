import docx
from styles import add_styles
from data import data

doc = docx.Document("./question_paper_template.docx")
add_styles(doc)


doc.add_paragraph(data["title"], "title")
doc.add_paragraph(f"Academic Year: {data["academic_year"]}", "para_center")

table = doc.add_table(rows=1, cols=2, style="table")

# Access the first and second columns (cells)
cell1 = table.cell(0, 0)
cell2 = table.cell(0, 1)

# Add text to the first column (cell)
cell1.text = 'Department: \nCourse: \nDate:\nMarks:'

# Add text to the second column (cell)
cell2.text = 'Year: \tSemester: \nTime: \nNo. of Pages: '

# Save the document
doc.save("Question.docx")
